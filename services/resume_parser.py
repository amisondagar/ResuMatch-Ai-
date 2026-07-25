"""
services/resume_parser.py — Resume File Parser
ResumeMatch.ai

Handles reading PDF and DOCX files and extracting raw text from them.

SUPPORTED FORMATS:
  .pdf  → pdfplumber (primary), PyPDF2 (fallback)
  .docx → python-docx

HOW IT WORKS:
  parse_resume(filepath, file_type)
    → reads the file
    → extracts raw text
    → calls ai.extractor to parse fields
    → returns (raw_text, parsed_data)

INTERVIEW TIP:
"We use pdfplumber for PDFs because it handles tables, columns,
and multi-line text better than PyPDF2."
"""

import os


def extract_text_from_pdf(filepath):
    """Extract text from a PDF file using pdfplumber, with PyPDF2 fallback."""
    text = ''

    # Primary: pdfplumber
    try:
        import pdfplumber
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + '\n'
        if text.strip():
            return text
    except Exception as e:
        print(f"pdfplumber failed: {e}")

    # Fallback: PyPDF2
    try:
        import PyPDF2
        with open(filepath, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + '\n'
        if text.strip():
            return text
    except Exception as e:
        print(f"PyPDF2 also failed: {e}")

    # Fallback 2: Direct text read if file is plain text saved with .pdf extension
    try:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            t = f.read()
            if t and len(t) > 10 and not t.startswith('%PDF'):
                return t
    except Exception:
        pass

    return text


def extract_text_from_docx(filepath):
    """Extract text from a DOCX file using python-docx."""
    text = ''
    try:
        import docx
        doc = docx.Document(filepath)
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        text = '\n'.join(paragraphs)

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text for cell in row.cells if cell.text.strip())
                if row_text:
                    text += '\n' + row_text
    except Exception as e:
        print(f"python-docx failed: {e}")

    return text


def parse_resume(filepath, file_type, original_filename=''):
    """
    Full parse pipeline:
      1. Extract raw text based on file type
      2. Run NLP extractor to get structured fields
      3. Return (raw_text, parsed_data)

    Args:
        filepath          (str): Absolute path to the uploaded file
        file_type         (str): 'pdf' or 'docx'
        original_filename (str): Original filename for name parsing fallback

    Returns:
        tuple: (raw_text: str, parsed_data: dict)
    """
    from ai.extractor import extractor
    from utils.helpers import clean_text

    # Step 1: Extract raw text
    ext = file_type.lower().lstrip('.')
    if ext == 'pdf':
        raw_text = extract_text_from_pdf(filepath)
    elif ext in ('docx', 'doc'):
        raw_text = extract_text_from_docx(filepath)
    else:
        raw_text = ''

    raw_text = clean_text(raw_text)

    # Step 2: Run NLP extractor
    if raw_text.strip():
        parsed_data = extractor.extract(raw_text, original_filename)
    else:
        parsed_data = {}

    return raw_text, parsed_data
